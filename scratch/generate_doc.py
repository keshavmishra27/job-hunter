import docx
from docx.shared import Pt, Inches

def create_doc():
    doc = docx.Document()

           
    title = doc.add_heading('How Job Hunter Scrapes Internships', 0)
    title.alignment = 1          

                 
    doc.add_heading('1. Overview', level=1)
    doc.add_paragraph(
        "The platform aggregates internship opportunities from multiple sources across the web. "
        "This is primarily done using Python and the asynchronous HTTP library 'httpx' combined with HTML parsing via 'BeautifulSoup'. "
        "The fetchers run asynchronously, allowing the system to scrape multiple sources concurrently for high performance."
    )

                        
    doc.add_heading('2. Scraping Indeed', level=1)
    doc.add_paragraph("Indeed is scraped using a custom IndeedFetcher class.")
    
    p = doc.add_paragraph()
    p.add_run('Search Query: ').bold = True
    p.add_run("It constructs a search URL with the keywords (e.g., 'python OR django') and location. It limits the search to the past 7 days (fromage=7) to ensure freshness.")
    p.style = 'List Bullet'
    
    p = doc.add_paragraph()
    p.add_run('Parsing the Cards: ').bold = True
    p.add_run("The search results page contains job cards (div.job_seen_beacon). The script extracts the job title, company name, location, and apply link using CSS selectors.")
    p.style = 'List Bullet'
    
    p = doc.add_paragraph()
    p.add_run('Enrichment and Expiration Check: ').bold = True
    p.add_run("For each card, the fetcher visits the detailed job page concurrently (max 5 at a time) to extract the full job description. It also checks for expired listings by looking for phrases like 'this job has expired' or 'this job is no longer available'. If an expiration phrase is found, the job is discarded.")
    p.style = 'List Bullet'

                             
    doc.add_heading('3. Scraping Internshala', level=1)
    doc.add_paragraph("Internshala is scraped using the InternshalaFetcher class.")
    
    p = doc.add_paragraph()
    p.add_run('URL Construction: ').bold = True
    p.add_run("The URL is constructed based on the location. If the user searches for 'remote', it points to internships/work-from-home-keywords-...")
    p.style = 'List Bullet'
    
    p = doc.add_paragraph()
    p.add_run('Data Extraction: ').bold = True
    p.add_run("The scraper targets the '.individual_internship' cards on the search results page. It extracts the title, company, location, stipend, and duration.")
    p.style = 'List Bullet'
    
    p = doc.add_paragraph()
    p.add_run('Remote Detection: ').bold = True
    p.add_run("It smartly infers if a job is remote or hybrid by analyzing the location string, searching for 'work from home', 'remote', or 'hybrid' keywords.")
    p.style = 'List Bullet'

                              
    doc.add_heading('4. Other Scraped Portals', level=1)
    doc.add_paragraph("The platform is also built to extract from other portals like:")
    
    p = doc.add_paragraph()
    p.add_run('LinkedIn: ').bold = True
    p.add_run("Fetched via public job search without authentication. It extracts listings from the base-card elements and infers remote modes.")
    p.style = 'List Bullet'
    
    p = doc.add_paragraph()
    p.add_run('Foundit, Freshersworld & Cutshort: ').bold = True
    p.add_run("Uses similar request and BeautifulSoup parsing mechanisms to scrape job titles, companies, locations, and descriptions.")
    p.style = 'List Bullet'

                                  
    doc.add_heading('5. Telegram Channel Scraping', level=1)
    doc.add_paragraph("The Telegram scraper (TelegramChannelFetcher) is highly unique. It does not require a Bot token or an API key, meaning it operates entirely independently of the Telegram API.")
    
    p = doc.add_paragraph()
    p.add_run('Web Preview: ').bold = True
    p.add_run("It utilizes Telegram's public web preview feature (t.me/s/<channel_username>).")
    p.style = 'List Bullet'
    
    p = doc.add_paragraph()
    p.add_run('Channel List: ').bold = True
    p.add_run("It targets a curated list of active public Telegram channels (e.g., JobsAndInternshipsIndia, internshipsalert).")
    p.style = 'List Bullet'
    
    p = doc.add_paragraph()
    p.add_run('Message Filtering: ').bold = True
    p.add_run("It downloads the HTML of the channel's web preview and parses the messages (.tgme_widget_message). It filters out non-internship messages by checking if the message text contains keywords like 'intern', 'internship', 'fresher', 'apply now', etc.")
    p.style = 'List Bullet'

    p = doc.add_paragraph()
    p.add_run('Link Extraction: ').bold = True
    p.add_run("It intelligently extracts the application link from the message by looking for links that contain text like 'apply', 'form', 'careers', or 'jobs'. It filters out internal telegram links (tg://).")
    p.style = 'List Bullet'

    p = doc.add_paragraph()
    p.add_run('Entity Extraction: ').bold = True
    p.add_run("The role title is inferred from the first few lines of the text. The company name is extracted using Regex patterns like '@CompanyName' or 'at XYZ'. The location is also inferred by searching for common city names.")
    p.style = 'List Bullet'

    p = doc.add_paragraph()
    p.add_run('Distribution: ').bold = True
    p.add_run("Once these notices are parsed, the telegram_sender.py module can optionally forward these eligible notices directly to a user's personal Telegram chat via the official Telegram Bot API (if they have configured their Chat ID).")
    p.style = 'List Bullet'

    doc.save('d:/kfiles/job-hunter/Scraping_Documentation.docx')
    print("Document saved to d:/kfiles/job-hunter/Scraping_Documentation.docx")

if __name__ == "__main__":
    create_doc()
